from __future__ import annotations

from pathlib import Path
from typing import Any

from docx import Document
from docx.document import Document as DocxDocument
from docx.table import Table

from app.rag.extractors.base import (
    BaseDocumentExtractor,
    DocumentExtractionError,
    EmptyDocumentError,
    ExtractedSection,
    ExtractionResult,
)


class DocxDocumentExtractor(BaseDocumentExtractor):
    """
    Extract structured text from Microsoft Word DOCX documents.

    Paragraphs are grouped by headings so that document structure,
    engineering sections, and future citations remain traceable.
    """

    supported_extensions = frozenset({".docx"})

    def extract(self, file_path: Path) -> ExtractionResult:
        path = self.validate_file(file_path)

        try:
            document = Document(path)
        except Exception as exc:
            raise DocumentExtractionError(
                f"Failed to open DOCX document '{path.name}'."
            ) from exc

        try:
            sections = self._extract_sections(document)

            if not sections:
                raise EmptyDocumentError(
                    f"No readable text was found in DOCX document '{path.name}'."
                )

            return ExtractionResult(
                source_path=path,
                file_type="docx",
                extraction_method="PythonDocxStructuredExtractor",
                sections=tuple(sections),
                metadata=self._build_document_metadata(document),
            )

        except DocumentExtractionError:
            raise

        except Exception as exc:
            raise DocumentExtractionError(
                f"Unexpected DOCX extraction failure for '{path.name}'."
            ) from exc

    def _extract_sections(
        self,
        document: DocxDocument,
    ) -> list[ExtractedSection]:
        """
        Group paragraphs and tables under their nearest heading.

        Heading paragraphs start a new section. Content before the first
        heading is preserved as an introductory section.
        """

        sections: list[ExtractedSection] = []
        current_heading: str | None = None
        current_heading_level: int | None = None
        current_content: list[str] = []
        section_index = 0

        def flush_section() -> None:
            nonlocal current_content, section_index

            section_text = "\n\n".join(
                item.strip()
                for item in current_content
                if item.strip()
            ).strip()

            if not section_text:
                current_content = []
                return

            sections.append(
                ExtractedSection(
                    text=section_text,
                    heading=current_heading,
                    metadata={
                        "section_index": section_index,
                        "heading_level": current_heading_level,
                        "section_type": (
                            "heading_section"
                            if current_heading is not None
                            else "introduction"
                        ),
                    },
                )
            )

            section_index += 1
            current_content = []

        for block in self._iter_document_blocks(document):
            if isinstance(block, str):
                text = block.strip()

                if not text:
                    continue

                heading_level = self._get_heading_level(block, document)

                if heading_level is not None:
                    flush_section()

                    current_heading = text
                    current_heading_level = heading_level
                    current_content = [text]
                else:
                    current_content.append(text)

            elif isinstance(block, Table):
                table_text = self._extract_table_text(block)

                if table_text:
                    current_content.append(table_text)

        flush_section()

        return sections

    @staticmethod
    def _iter_document_blocks(
        document: DocxDocument,
    ) -> list[str | Table]:
        """
        Return paragraphs and tables in document order.

        python-docx exposes paragraphs and tables separately by default,
        so XML traversal is used to preserve their original ordering.
        """

        blocks: list[str | Table] = []

        for child in document.element.body.iterchildren():
            tag_name = child.tag.rsplit("}", 1)[-1]

            if tag_name == "p":
                paragraph = next(
                    (
                        paragraph
                        for paragraph in document.paragraphs
                        if paragraph._element is child
                    ),
                    None,
                )

                if paragraph is not None:
                    blocks.append(paragraph.text)

            elif tag_name == "tbl":
                table = next(
                    (
                        table
                        for table in document.tables
                        if table._element is child
                    ),
                    None,
                )

                if table is not None:
                    blocks.append(table)

        return blocks

    @staticmethod
    def _get_heading_level(
        paragraph_text: str,
        document: DocxDocument,
    ) -> int | None:
        """
        Resolve heading level by matching paragraph text and style.

        Returns None when the paragraph is not a Word heading.
        """

        normalized_text = paragraph_text.strip()

        for paragraph in document.paragraphs:
            if paragraph.text.strip() != normalized_text:
                continue

            style_name = paragraph.style.name if paragraph.style else ""

            if not style_name.lower().startswith("heading"):
                return None

            suffix = style_name.replace("Heading", "").strip()

            try:
                return int(suffix)
            except ValueError:
                return 1

        return None

    @staticmethod
    def _extract_table_text(table: Table) -> str:
        """
        Convert a DOCX table into readable pipe-separated text.

        Empty rows are ignored. This preserves table content for later
        chunking and semantic retrieval without requiring a separate format.
        """

        rows: list[str] = []

        for row in table.rows:
            cells = [
                cell.text.strip().replace("\n", " ")
                for cell in row.cells
            ]

            if not any(cells):
                continue

            rows.append(" | ".join(cells))

        return "\n".join(rows)

    @staticmethod
    def _build_document_metadata(
        document: DocxDocument,
    ) -> dict[str, Any]:
        """Extract safe DOCX-level metadata."""

        properties = document.core_properties

        metadata: dict[str, Any] = {
            "paragraph_count": len(document.paragraphs),
            "table_count": len(document.tables),
            "format": "DOCX",
        }

        optional_properties = {
            "title": properties.title,
            "subject": properties.subject,
            "author": properties.author,
            "category": properties.category,
            "comments": properties.comments,
            "keywords": properties.keywords,
            "last_modified_by": properties.last_modified_by,
        }

        metadata["document_properties"] = {
            key: value.strip()
            for key, value in optional_properties.items()
            if isinstance(value, str) and value.strip()
        }

        return metadata